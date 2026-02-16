import os
import uuid

from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.oxml.ns import qn

TEMP_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "temp")

GRAY_COLOR = RGBColor(0xB0, 0xB0, 0xB0)


def _set_run_font(run, font_name: str, size: int, east_asia_font: str | None = None, color: RGBColor | None = None):
    """Set font properties on a run."""
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = run._r.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    ea = east_asia_font or font_name
    rFonts.set(qn("w:eastAsia"), ea)


def _set_cell_width(cell, width_emu: int):
    """Set cell width in EMU via XML."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = cell._tc.makeelement(qn("w:tcW"), {})
        tcPr.append(tcW)
    # Convert EMU to twips (1 twip = 635 EMU)
    twips = int(width_emu / 635)
    tcW.set(qn("w:w"), str(twips))
    tcW.set(qn("w:type"), "dxa")


def _set_cell_no_wrap(cell):
    """Prevent text wrapping in a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    noWrap = tcPr.find(qn("w:noWrap"))
    if noWrap is None:
        noWrap = cell._tc.makeelement(qn("w:noWrap"), {})
        tcPr.append(noWrap)


def _calc_chinese_text_width_pt(text: str, font_size_pt: int) -> float:
    """Estimate text width in points for Chinese characters.

    CJK characters ~= font_size width each.
    ASCII/punctuation ~= font_size * 0.5 width each.
    """
    width = 0.0
    for ch in text:
        if ord(ch) > 0x7F:
            width += font_size_pt
        else:
            width += font_size_pt * 0.5
    return width


def generate_workbook_type1(entries: list[dict], job_id: str | None = None) -> str:
    """Generate Type 1 workbook: table with gray Chinese characters for tracing.

    Column widths are adjusted so 한자 and 예문 columns fit in one line.
    병음 and 의미 columns may wrap to 2 lines.
    """
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # Available page width = A4 width (21cm) - left margin - right margin
    page_width_cm = 21.0 - 2.0 - 2.0  # 17cm
    page_width_pt = page_width_cm / 2.54 * 72  # ~481pt

    # Calculate required widths for 한자 and 예문 columns
    chinese_font_size = 16
    max_chinese_width_pt = 0.0
    max_example_width_pt = 0.0

    for entry in entries:
        chinese_w = _calc_chinese_text_width_pt(entry.get("chinese", ""), chinese_font_size)
        example_w = _calc_chinese_text_width_pt(entry.get("example", ""), chinese_font_size)
        max_chinese_width_pt = max(max_chinese_width_pt, chinese_w)
        max_example_width_pt = max(max_example_width_pt, example_w)

    # Add padding (cell margins ~10pt each side)
    cell_padding = 20
    chinese_col_pt = max_chinese_width_pt + cell_padding
    example_col_pt = max_example_width_pt + cell_padding

    # Remaining width split equally between 병음 and 의미
    remaining_pt = page_width_pt - chinese_col_pt - example_col_pt
    min_col_pt = 50  # minimum column width
    if remaining_pt < min_col_pt * 2:
        remaining_pt = min_col_pt * 2
        # Recalculate: shrink Chinese columns proportionally
        total_chinese = chinese_col_pt + example_col_pt
        available_for_chinese = page_width_pt - remaining_pt
        ratio = available_for_chinese / total_chinese if total_chinese > 0 else 1
        chinese_col_pt = chinese_col_pt * ratio
        example_col_pt = example_col_pt * ratio

    pinyin_col_pt = remaining_pt / 2
    meaning_col_pt = remaining_pt / 2

    # Convert pt to EMU (1 pt = 12700 EMU)
    pt_to_emu = 12700
    chinese_col_emu = int(chinese_col_pt * pt_to_emu)
    pinyin_col_emu = int(pinyin_col_pt * pt_to_emu)
    meaning_col_emu = int(meaning_col_pt * pt_to_emu)
    example_col_emu = int(example_col_pt * pt_to_emu)

    # Create table
    table = doc.add_table(rows=len(entries), cols=4)
    table.style = 'Table Grid'
    table.autofit = False

    for i, entry in enumerate(entries):
        row = table.rows[i]
        cells = row.cells

        # Set column widths
        _set_cell_width(cells[0], chinese_col_emu)
        _set_cell_width(cells[1], pinyin_col_emu)
        _set_cell_width(cells[2], meaning_col_emu)
        _set_cell_width(cells[3], example_col_emu)

        # Prevent wrapping on Chinese columns
        _set_cell_no_wrap(cells[0])
        _set_cell_no_wrap(cells[3])

        # 한자 - Gray color for tracing
        p = cells[0].paragraphs[0]
        p.clear()
        run = p.add_run(entry.get("chinese", ""))
        _set_run_font(run, "Microsoft YaHei", 16, color=GRAY_COLOR)

        # 병음
        p = cells[1].paragraphs[0]
        p.clear()
        run = p.add_run(entry.get("pinyin", ""))
        _set_run_font(run, "맑은 고딕", 11)

        # 의미
        p = cells[2].paragraphs[0]
        p.clear()
        run = p.add_run(entry.get("meaning", ""))
        _set_run_font(run, "맑은 고딕", 11)

        # 예문 - Gray color for tracing (same as 한자)
        p = cells[3].paragraphs[0]
        p.clear()
        example_text = entry.get("example", "")
        if example_text:
            run = p.add_run(example_text)
            _set_run_font(run, "Microsoft YaHei", 16, color=GRAY_COLOR)

    # Save
    os.makedirs(TEMP_DIR, exist_ok=True)
    file_id = job_id or str(uuid.uuid4())
    output_path = os.path.join(TEMP_DIR, f"{file_id}.docx")
    doc.save(output_path)
    return output_path


def generate_workbook_type2(entries: list[dict], job_id: str | None = None) -> str:
    """Generate Type 2 workbook: Korean interpretation then Chinese text (gray) for tracing."""
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    for entry in entries:
        speaker = entry.get("speaker", "")
        korean = entry.get("korean", "")
        chinese_text = entry.get("chinese_text", "")

        prefix = f"{speaker}: " if speaker else ""

        # Korean interpretation line - no extra spacing
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(prefix + korean)
        _set_run_font(run, "맑은 고딕", 11)

        # Chinese text line - Gray for tracing, no extra spacing
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(prefix + chinese_text)
        _set_run_font(run, "Microsoft YaHei", 16, color=GRAY_COLOR)

    # Save
    os.makedirs(TEMP_DIR, exist_ok=True)
    file_id = job_id or str(uuid.uuid4())
    output_path = os.path.join(TEMP_DIR, f"{file_id}.docx")
    doc.save(output_path)
    return output_path


def generate_workbook(entries: list[dict], workbook_type: str, job_id: str | None = None) -> str:
    if workbook_type == "type1":
        return generate_workbook_type1(entries, job_id)
    return generate_workbook_type2(entries, job_id)
