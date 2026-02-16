import copy
import os
import uuid

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "..", "..", "중국어단어장.docx")
TEMP_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "temp")

MIN_FONT_SIZE = 7


def _get_cell_width_pt(cell) -> float:
    """Read cell width from XML (dxa -> pt)."""
    tcPr = cell._tc.find(qn("w:tcPr"))
    if tcPr is not None:
        tcW = tcPr.find(qn("w:tcW"))
        if tcW is not None:
            return int(tcW.get(qn("w:w"))) / 20.0
    return 90.0


def _get_cell_padding_pt(table, cell) -> float:
    """Read left+right cell padding from cell -> table -> style (dxa -> pt)."""
    # 1) Cell-level margins
    tcPr = cell._tc.find(qn("w:tcPr"))
    if tcPr is not None:
        tcMar = tcPr.find(qn("w:tcMar"))
        if tcMar is not None:
            return _sum_lr_margin(tcMar)

    # 2) Table-level margins
    tblPr = table._tbl.find(qn("w:tblPr"))
    if tblPr is not None:
        tblCellMar = tblPr.find(qn("w:tblCellMar"))
        if tblCellMar is not None:
            return _sum_lr_margin(tblCellMar)

    # 3) Style-level (Normal Table default)
    doc_element = table._tbl.getroottree().getroot()
    for style_el in doc_element.iter(qn("w:style")):
        style_tblPr = style_el.find(qn("w:tblPr"))
        if style_tblPr is not None:
            mar = style_tblPr.find(qn("w:tblCellMar"))
            if mar is not None:
                return _sum_lr_margin(mar)

    return 10.8  # Word default fallback


def _sum_lr_margin(margin_element) -> float:
    """Sum left and right margins from a tblCellMar or tcMar element (dxa -> pt)."""
    total_dxa = 0
    for side in ("w:left", "w:start", "w:right", "w:end"):
        el = margin_element.find(qn(side))
        if el is not None:
            total_dxa += int(el.get(qn("w:w"), "0"))
    return total_dxa / 20.0


def _calc_font_size(text: str, usable_width_pt: float, base_size: int) -> int:
    """Calculate font size to fit text in one line.

    Full-width chars (Korean, Chinese) count as 1.0 width unit.
    Half-width chars (ASCII, spaces, punctuation) count as 0.5 width unit.
    """
    width_units = 0
    for ch in text:
        if ord(ch) > 0x7F:
            width_units += 1.0
        else:
            width_units += 0.5

    if width_units == 0:
        return base_size

    max_size = usable_width_pt / width_units
    font_size = min(base_size, int(max_size))
    return max(font_size, MIN_FONT_SIZE)


def _set_cell_text(cell, text: str, font_name: str, font_size: int, east_asia_font: str | None = None):
    """Set cell text with specified font and size."""
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(text)
    run.font.size = Pt(font_size)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = run._r.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    ea = east_asia_font or font_name
    rFonts.set(qn("w:eastAsia"), ea)


def generate_word(words: list[dict], job_id: str | None = None) -> str:
    """Generate a Word file from extracted words using the template."""
    doc = Document(os.path.abspath(TEMPLATE_PATH))
    body = doc.element.body
    table = doc.tables[0]

    # Read template dimensions dynamically
    rows_per_page = len(table.rows)
    korean_cell = table.rows[0].cells[5]
    cell_width = _get_cell_width_pt(korean_cell)
    cell_padding = _get_cell_padding_pt(table, korean_cell)
    usable_width = cell_width - cell_padding

    # Save original table XML for cloning
    template_tbl_xml = copy.deepcopy(table._tbl)

    # Split words into pages
    pages = []
    for i in range(0, len(words), rows_per_page):
        pages.append(words[i : i + rows_per_page])

    if not pages:
        pages = [[]]

    # Fill first page (existing table)
    _fill_table(doc.tables[0], pages[0], usable_width)

    # Add additional pages by cloning the template table
    for page_idx in range(1, len(pages)):
        new_tbl = copy.deepcopy(template_tbl_xml)
        body.append(new_tbl)
        new_table = doc.tables[page_idx]
        _fill_table(new_table, pages[page_idx], usable_width)

    # Save to temp directory
    os.makedirs(TEMP_DIR, exist_ok=True)
    file_id = job_id or str(uuid.uuid4())
    output_path = os.path.join(TEMP_DIR, f"{file_id}.docx")
    doc.save(output_path)

    return output_path


def _fill_table(table, words: list[dict], usable_width: float):
    """Fill a table's rows with word data."""
    for i, word in enumerate(words):
        if i >= len(table.rows):
            break
        row = table.rows[i]
        cells = row.cells
        _set_cell_text(cells[0], word.get("chinese", ""), "Microsoft YaHei", 12)
        _set_cell_text(cells[3], word.get("pinyin", ""), "맑은 고딕", 12)
        korean_text = word.get("korean", "")
        korean_size = _calc_font_size(korean_text, usable_width, 11)
        _set_cell_text(cells[5], korean_text, "맑은 고딕", korean_size)
